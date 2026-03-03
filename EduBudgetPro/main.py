"""
╔══════════════════════════════════════════════════════════════════════╗
║         EduBudget Pro v3 – Student Learning & Budget System           ║
║   • Smart Budget with 50/30/20 Rule + Category Checkboxes             ║
║   • PDF/Word → Offline Keyword-Based Questions (no API key needed)    ║
║   • Optional Claude AI questions (needs API key)                      ║
║   • Fully fixed & clean codebase                                      ║
╚══════════════════════════════════════════════════════════════════════╝

INSTALL (one-time):
    pip install pypdf python-docx matplotlib requests

RUN:
    python edubudget_pro_v3.py
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import json, os, random, datetime, csv, re, threading, math

# ── Optional imports ────────────────────────────────────────────────────────
try:
    import matplotlib; matplotlib.use("TkAgg")
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        HAS_PDF = True
    except ImportError:
        HAS_PDF = False

try:
    import docx as _docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import requests as _req
    HAS_REQ = True
except ImportError:
    HAS_REQ = False

# ═══════════════════════════════════════════════════════════════════════
#  THEME
# ═══════════════════════════════════════════════════════════════════════
BG       = "#0d1117"
CARD     = "#161b22"
INPUT    = "#21262d"
BORDER   = "#30363d"
ACCENT   = "#58a6ff"
PURPLE   = "#bc8cff"
GREEN    = "#3fb950"
RED      = "#f85149"
ORANGE   = "#f0883e"
YELLOW   = "#d29922"
TEXT     = "#c9d1d9"
DIM      = "#8b949e"
BRIGHT   = "#e6edf3"
SEL      = "#1f6feb"
AI_CLR   = "#f78166"

F        = ("Segoe UI", 10)
FB       = ("Segoe UI", 10, "bold")
FS       = ("Segoe UI", 9)
FH       = ("Segoe UI", 14, "bold")
FT       = ("Segoe UI", 20, "bold")
FM       = ("Segoe UI", 12, "bold")

DATA_FILE = "edubudget_v3.json"
API_URL   = "https://api.anthropic.com/v1/messages"
API_MODEL = "claude-haiku-4-5-20251001"   # cheaper/faster model

# ═══════════════════════════════════════════════════════════════════════
#  DATA
# ═══════════════════════════════════════════════════════════════════════
ALL_CATEGORIES = [
    "Food", "Rent", "Transport", "Shopping",
    "Entertainment", "Utilities", "Healthcare",
    "Education", "Savings", "Other"
]

def default_data():
    return {
        "exams_generated": 0,
        "api_key": "",
        "history": {},
        "budget": {
            "month": datetime.datetime.now().strftime("%Y-%m"),
            "income": 0,
            "savings_goal": 0,
            "active_categories": ["Food", "Rent", "Transport", "Shopping", "Utilities"],
            "expenses": {},          # cat -> amount
            "custom_expenses": [],   # [{name, amount, category}]
        }
    }

def load_data():
    if os.path.exists(DATA_FILE):
        try:
            with open(DATA_FILE) as f:
                d = json.load(f)
            def_d = default_data()
            for k, v in def_d.items():
                if k not in d:
                    d[k] = v
            for k, v in def_d["budget"].items():
                if k not in d.get("budget", {}):
                    d.setdefault("budget", {})[k] = v
            return d
        except Exception:
            pass
    return default_data()

def save_data(d):
    try:
        with open(DATA_FILE, "w") as f:
            json.dump(d, f, indent=2)
    except Exception as e:
        messagebox.showerror("Save Error", str(e))

# ═══════════════════════════════════════════════════════════════════════
#  UI HELPERS
# ═══════════════════════════════════════════════════════════════════════
def btn(parent, text, cmd=None, bg=ACCENT, fg=BG, font=FB,
        padx=12, pady=6, **kw):
    return tk.Button(parent, text=text, command=cmd, bg=bg, fg=fg,
                     font=font, relief="flat", bd=0, cursor="hand2",
                     activebackground=SEL, activeforeground="white",
                     padx=padx, pady=pady, **kw)

def entry(parent, var=None, w=18, show=None):
    kw = dict(bg=INPUT, fg=BRIGHT, insertbackground=BRIGHT,
               relief="flat", font=F, bd=4)
    if show: kw["show"] = show
    return tk.Entry(parent, textvariable=var, width=w, **kw)

def spin(parent, var, lo=0, hi=9999, w=9):
    return tk.Spinbox(parent, from_=lo, to=hi, textvariable=var,
                      width=w, bg=INPUT, fg=BRIGHT, relief="flat",
                      font=F, bd=3, buttonbackground=CARD,
                      insertbackground=BRIGHT)

def lbl(parent, text="", font=F, fg=TEXT, bg=CARD, **kw):
    return tk.Label(parent, text=text, font=font, fg=fg, bg=bg, **kw)

def sep(parent, color=BORDER, pady=6):
    tk.Frame(parent, bg=color, height=1).pack(fill="x", pady=pady)

def scrollable(parent, bg=CARD):
    c = tk.Canvas(parent, bg=bg, highlightthickness=0)
    vs = ttk.Scrollbar(parent, orient="vertical", command=c.yview)
    c.configure(yscrollcommand=vs.set)
    vs.pack(side="right", fill="y")
    c.pack(side="left", fill="both", expand=True)
    inner = tk.Frame(c, bg=bg)
    wid = c.create_window((0, 0), window=inner, anchor="nw")

    def _resize(e):
        c.configure(scrollregion=c.bbox("all"))
        c.itemconfig(wid, width=c.winfo_width())

    inner.bind("<Configure>", _resize)
    c.bind("<Configure>", lambda e: c.itemconfig(wid, width=c.winfo_width()))

    def _scroll(e):
        c.yview_scroll(int(-1 * (e.delta / 120)), "units")
    c.bind_all("<MouseWheel>", _scroll)
    return inner

# ═══════════════════════════════════════════════════════════════════════
#  BUDGET LOGIC  
# ═══════════════════════════════════════════════════════════════════════
def smart_split(income, savings_goal, active_cats):
    """
    50/30/20 inspired auto-split across the user's chosen categories.
    Needs   = Rent, Food, Utilities, Transport, Healthcare, Education
    Wants   = Shopping, Entertainment, Other
    Savings = already set by user
    """
    spendable = max(0, income - savings_goal)
    if not active_cats or spendable == 0:
        return {c: 0 for c in active_cats}

    NEED_CATS  = {"Rent", "Food", "Utilities", "Transport", "Healthcare", "Education"}
    WANT_CATS  = {"Shopping", "Entertainment", "Other"}

    needs = [c for c in active_cats if c in NEED_CATS]
    wants = [c for c in active_cats if c in WANT_CATS]
    # Savings handled externally
    has_savings_cat = "Savings" in active_cats

    # proportional budgets
    need_pool  = spendable * 0.60
    want_pool  = spendable * 0.30
    sav_pool   = spendable * 0.10

    # default weights for needs
    NEED_WEIGHTS = {
        "Rent": 4, "Food": 3, "Utilities": 1.5,
        "Transport": 1.5, "Healthcare": 1, "Education": 1
    }
    WANT_WEIGHTS = {"Shopping": 2, "Entertainment": 1.5, "Other": 1}

    result = {}

    if needs:
        total_w = sum(NEED_WEIGHTS.get(c, 1) for c in needs)
        for c in needs:
            result[c] = round(need_pool * NEED_WEIGHTS.get(c, 1) / total_w, 2)
    else:
        want_pool += need_pool

    if wants:
        total_w = sum(WANT_WEIGHTS.get(c, 1) for c in wants)
        for c in wants:
            result[c] = round(want_pool * WANT_WEIGHTS.get(c, 1) / total_w, 2)

    if has_savings_cat:
        result["Savings"] = round(sav_pool, 2)

    return result

def calc_budget(budget_data):
    income       = float(budget_data.get("income", 0))
    savings_goal = float(budget_data.get("savings_goal", 0))
    expenses     = {k: float(v) for k, v in budget_data.get("expenses", {}).items()}
    customs      = budget_data.get("custom_expenses", [])

    cat_totals = dict(expenses)
    for exp in customs:
        cat = exp.get("category", "Other")
        cat_totals[cat] = cat_totals.get(cat, 0) + float(exp.get("amount", 0))

    total_exp  = sum(cat_totals.values())
    remaining  = income - savings_goal - total_exp
    top_cat    = max(cat_totals, key=cat_totals.get) if cat_totals else "N/A"
    top_amt    = cat_totals.get(top_cat, 0)

    return {
        "income": income, "savings_goal": savings_goal,
        "cat_totals": cat_totals, "total_exp": total_exp,
        "remaining": remaining, "top_cat": top_cat,
        "top_amt": top_amt, "over": remaining < 0,
        "spendable": income - savings_goal,
    }

# ═══════════════════════════════════════════════════════════════════════
#  PDF / WORD TEXT EXTRACTION
# ═══════════════════════════════════════════════════════════════════════
def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        if not HAS_PDF:
            raise RuntimeError("Install pypdf:  pip install pypdf")
        reader = PdfReader(path)
        pages  = [p.extract_text() or "" for p in reader.pages]
        text   = "\n\n".join(p.strip() for p in pages if p.strip())
        if not text:
            raise ValueError("No readable text in PDF (may be scanned image).")
        return text
    elif ext in (".docx", ".doc"):
        if not HAS_DOCX:
            raise RuntimeError("Install python-docx:  pip install python-docx")
        doc  = _docx.Document(path)
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paras.append(cell.text.strip())
        text = "\n\n".join(paras)
        if not text:
            raise ValueError("No text found in Word document.")
        return text
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use .pdf or .docx")

# ═══════════════════════════════════════════════════════════════════════
#  OFFLINE QUESTION GENERATOR (no API needed)
# ═══════════════════════════════════════════════════════════════════════
def _sentences(text):
    sents = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sents if len(s.strip()) > 30]

def _key_terms(text, n=20):
    """Extract important noun-like terms from text."""
    words = re.findall(r'\b[A-Z][a-z]{3,}\b|\b[a-z]{5,}\b', text)
    stop  = {"about","their","there","which","where","while","these",
              "those","would","could","should","other","using","after",
              "before","being","having","since","under","above","below",
              "when","with","from","this","that","also","into","over"}
    freq  = {}
    for w in words:
        w = w.lower()
        if w not in stop:
            freq[w] = freq.get(w, 0) + 1
    return [k for k, _ in sorted(freq.items(), key=lambda x: -x[1])[:n]]

def _make_distractors(correct, all_terms):
    """Pick 3 plausible wrong answers."""
    pool = [t.capitalize() for t in all_terms
            if t.lower() != correct.lower()]
    random.shuffle(pool)
    return pool[:3]

def offline_generate_questions(text, config):
    """
    Generate exam questions entirely offline using the document text.
    Returns a dict with the same structure as the AI response.
    """
    n_easy   = config.get("easy", 3)
    n_medium = config.get("medium", 3)
    n_hard   = config.get("hard", 2)
    n_mcq    = config.get("mcq_count", 5)
    n_num    = config.get("numerical_count", 0)

    sents = _sentences(text)
    terms = _key_terms(text, 30)
    paras = [p.strip() for p in text.split("\n\n") if len(p.strip()) > 60]

    random.shuffle(sents)
    random.shuffle(paras)

    # ── MCQs ──────────────────────────────────────────────────────────
    mcqs = []
    pool = list(sents)
    random.shuffle(pool)
    for sent in pool:
        if len(mcqs) >= n_mcq:
            break
        # Find a key term in the sentence to blank out
        found = [(m.group(), m.start(), m.end())
                 for m in re.finditer(r'\b[A-Z][a-z]{3,}\b', sent)]
        if not found:
            found = [(m.group(), m.start(), m.end())
                     for m in re.finditer(r'\b[a-z]{5,}\b', sent)]
        if not found:
            continue
        term, s, e = random.choice(found)
        question = sent[:s] + "_______" + sent[e:]
        correct  = term.capitalize()
        distractors = _make_distractors(correct, terms)
        if len(distractors) < 3:
            continue
        options = [correct] + distractors[:3]
        random.shuffle(options)
        ci = options.index(correct)
        mcqs.append({
            "question": f"Fill in the blank: {question}",
            "options": options, "correct_index": ci,
            "correct_text": correct
        })

    # ── Easy: Definition/simple questions ─────────────────────────────
    easy_templates = [
        "Define the term '{}' as used in this topic.",
        "What is meant by '{}' in the context of this subject?",
        "Give a brief explanation of '{}'.",
        "In your own words, describe what '{}' means.",
        "What is '{}'? Give one example.",
    ]
    easy_qs = []
    t_pool = list(terms)
    random.shuffle(t_pool)
    for term in t_pool[:n_easy * 2]:
        if len(easy_qs) >= n_easy:
            break
        tmpl = random.choice(easy_templates)
        easy_qs.append(tmpl.format(term.capitalize()))

    # ── Medium: Explanation questions from paragraphs ─────────────────
    med_templates = [
        "Explain the concept described in the following excerpt and give a real-world example:\n  \"{}\"",
        "Describe and elaborate on this idea from the document:\n  \"{}\"",
        "Summarize and explain the significance of:\n  \"{}\"",
        "Analyze the following statement and explain its importance:\n  \"{}\"",
    ]
    medium_qs = []
    p_pool = [p[:180] + ("..." if len(p) > 180 else "") for p in paras]
    random.shuffle(p_pool)
    for para in p_pool[:n_medium * 2]:
        if len(medium_qs) >= n_medium:
            break
        tmpl = random.choice(med_templates)
        medium_qs.append(tmpl.format(para))

    # ── Hard: Analytical questions ────────────────────────────────────
    hard_templates = [
        "Critically analyze the topic covered in this document. What are its key strengths and limitations?",
        "How does the concept of '{}' connect to real-world applications? Justify your answer with examples from the document.",
        "Compare and contrast the main ideas presented. What conclusions can you draw?",
        "Evaluate the significance of '{}' as discussed in the material. Support your argument with evidence.",
        "Why is '{}' considered important in this field? Analyze its implications.",
        "Discuss how the ideas in this document could be applied to solve a practical problem.",
    ]
    hard_qs = []
    h_terms = terms[:10]
    random.shuffle(h_terms)
    for i in range(n_hard):
        tmpl = random.choice(hard_templates)
        t = h_terms[i % len(h_terms)] if h_terms else "this concept"
        hard_qs.append(tmpl.format(t.capitalize()))

    # ── Numericals: only if numbers found ─────────────────────────────
    num_qs = []
    if n_num > 0:
        num_sents = [s for s in sents
                     if re.search(r'\d+\.?\d*\s*(%|kg|km|m|s|hrs?|min|₹|\$)', s, re.I)]
        for ns in num_sents[:n_num]:
            num_qs.append({
                "question": f"Based on the data given: '{ns[:150]}' — solve or explain the numerical relationship.",
                "formula": "Refer to relevant formula from the document.",
                "solution_hint": "Extract values from the sentence and apply the appropriate formula."
            })
        while len(num_qs) < n_num:
            num_qs.append({
                "question": "Create a short numerical problem based on the key concepts in this document and solve it.",
                "formula": "Select appropriate formula from the topic.",
                "solution_hint": "Apply step-by-step approach from the document's principles."
            })

    # ── Answer key hints ───────────────────────────────────────────────
    easy_hints  = [f"Key term: {terms[i % len(terms)].capitalize() if terms else 'See document'}"
                   for i in range(len(easy_qs))]
    med_hints   = [f"Key points from paragraph {i+1} — explain in 3–5 sentences."
                   for i in range(len(medium_qs))]
    hard_hints  = [f"Analytical answer required — use document evidence. No single correct answer."
                   for _ in hard_qs]
    mcq_answers = [mcq["correct_text"] for mcq in mcqs]

    # Summarize document (first 200 chars of first paragraph)
    summary = paras[0][:200] + "..." if paras else text[:200] + "..."

    return {
        "subject_detected": config.get("subject", "Document-Based Subject"),
        "topic_summary": summary,
        "mcqs": mcqs[:n_mcq],
        "easy_questions": easy_qs[:n_easy],
        "medium_questions": medium_qs[:n_medium],
        "hard_questions": hard_qs[:n_hard],
        "numerical_problems": num_qs[:n_num],
        "answer_key": {
            "mcq_answers": mcq_answers,
            "easy_hints": easy_hints,
            "medium_hints": med_hints,
            "hard_hints": hard_hints,
        }
    }

# ═══════════════════════════════════════════════════════════════════════
#  AI QUESTION GENERATOR (Claude API)
# ═══════════════════════════════════════════════════════════════════════
def ai_generate_questions(doc_text, config, api_key, on_progress=None):
    if not HAS_REQ:
        raise RuntimeError("Install requests:  pip install requests")

    truncated = doc_text[:10000] + ("\n[...truncated...]" if len(doc_text) > 10000 else "")
    n_easy = config.get("easy", 3)
    n_med  = config.get("medium", 3)
    n_hard = config.get("hard", 2)
    n_mcq  = config.get("mcq_count", 5)
    n_num  = config.get("numerical_count", 0)
    marks  = config.get("total_marks", 100)
    subj   = config.get("subject", "the subject")

    prompt = f"""You are an expert exam creator. Read the document and generate exam questions ONLY from its content.

DOCUMENT:
---
{truncated}
---

REQUIREMENTS: {n_mcq} MCQs, {n_easy} easy, {n_med} medium, {n_hard} hard questions, {n_num} numericals. Total marks: {marks}. Subject hint: {subj}.

Respond ONLY with this exact JSON (no extra text):
{{
  "subject_detected": "...",
  "topic_summary": "2 sentence summary",
  "mcqs": [{{"question":"...","options":["A","B","C","D"],"correct_index":0,"correct_text":"A"}}],
  "easy_questions": ["..."],
  "medium_questions": ["..."],
  "hard_questions": ["..."],
  "numerical_problems": [{{"question":"...","formula":"...","solution_hint":"..."}}],
  "answer_key": {{
    "mcq_answers": ["correct answer text"],
    "easy_hints": ["key point"],
    "medium_hints": ["key points"],
    "hard_hints": ["analytical points"]
  }}
}}"""

    if on_progress:
        on_progress("Sending to Claude AI...")

    resp = _req.post(API_URL,
                     headers={"x-api-key": api_key,
                               "anthropic-version": "2023-06-01",
                               "content-type": "application/json"},
                     json={"model": API_MODEL, "max_tokens": 4096,
                           "messages": [{"role": "user", "content": prompt}]},
                     timeout=90)

    if resp.status_code == 401:
        raise ValueError("Invalid API key.")
    elif resp.status_code == 429:
        raise ValueError("Rate limit hit. Wait a moment and retry.")
    elif resp.status_code != 200:
        raise ValueError(f"API error {resp.status_code}: {resp.text[:200]}")

    raw = resp.json()["content"][0]["text"].strip()
    raw = raw.lstrip("```json").lstrip("```").rstrip("```").strip()
    try:
        return json.loads(raw)
    except Exception:
        s, e = raw.find("{"), raw.rfind("}") + 1
        return json.loads(raw[s:e])

# ═══════════════════════════════════════════════════════════════════════
#  FORMAT EXAM PAPER
# ═══════════════════════════════════════════════════════════════════════
def format_exam_paper(ai_data, config, source_label="AI-Generated"):
    W       = 78
    subj    = ai_data.get("subject_detected", config.get("subject", ""))
    title   = config.get("title", "Examination")
    school  = config.get("school_name", "")
    teacher = config.get("teacher_name", "")
    dur     = config.get("duration", "2 Hours")
    marks   = config.get("total_marks", 100)
    show_k  = config.get("show_answer_key", True)

    mcqs   = ai_data.get("mcqs", [])
    easy   = ai_data.get("easy_questions", [])
    med    = ai_data.get("medium_questions", [])
    hard   = ai_data.get("hard_questions", [])
    nums   = ai_data.get("numerical_problems", [])
    ak     = ai_data.get("answer_key", {})

    nm, ne, nmd, nh, nn = len(mcqs), len(easy), len(med), len(hard), len(nums)
    mm = 1
    em = max(2, round(marks * 0.10 / max(ne, 1)))
    mmd = max(3, round(marks * 0.25 / max(nmd, 1)))
    hm = max(5, round(marks * 0.30 / max(nh, 1)))
    nmm = max(4, round(marks * 0.25 / max(nn, 1)))

    L = []
    def p(s=""): L.append(s)
    def ctr(s): L.append(s.center(W))
    def rule(c="─"): L.append(c * W)

    rule("═")
    ctr(f"[{source_label}]")
    if school: ctr(school.upper())
    ctr(title.upper())
    ctr(f"Subject: {subj}")
    ctr(f"Duration: {dur}   |   Total Marks: {marks}")
    if teacher: ctr(f"Teacher: {teacher}")
    ctr(datetime.datetime.now().strftime("%d %B %Y"))
    rule("═")
    p()
    p("INSTRUCTIONS: All sections compulsory. Marks shown in [ ].")
    p("Show all working for calculations. Write legibly.")
    p()

    q = 1
    if mcqs:
        p(); ctr(f"SECTION A — MCQs  [{nm*mm} Marks]"); rule()
        p(f"  Choose the best option. [{mm} Mark each]\n")
        for m in mcqs:
            p(f"  Q{q}. {m['question']}  [{mm}]")
            for j, o in enumerate(m.get("options", [])):
                p(f"       ({chr(65+j)}) {o}")
            p(); q += 1

    if easy:
        p(); ctr(f"SECTION B — SHORT ANSWER  [{ne*em} Marks]"); rule()
        p(f"  Answer briefly. [{em} Marks each]\n")
        for s in easy:
            p(f"  Q{q}. {s}  [{em}]"); p(); q += 1

    if med:
        p(); ctr(f"SECTION C — DESCRIPTIVE  [{nmd*mmd} Marks]"); rule()
        p(f"  Write detailed answers. [{mmd} Marks each]\n")
        for s in med:
            p(f"  Q{q}. {s}  [{mmd}]"); p(); q += 1

    if hard:
        p(); ctr(f"SECTION D — ANALYTICAL  [{nh*hm} Marks]"); rule()
        p(f"  Justify with reasoning. [{hm} Marks each]\n")
        for s in hard:
            p(f"  Q{q}. {s}  [{hm}]"); p(); q += 1

    if nums:
        p(); ctr(f"SECTION E — NUMERICALS  [{nn*nmm} Marks]"); rule()
        p(f"  Show all steps. [{nmm} Marks each]\n")
        for n in nums:
            p(f"  Q{q}. {n.get('question','')}  [{nmm}]")
            if n.get("formula"):
                p(f"       (Formula: {n['formula']})")
            p(); q += 1

    actual = nm*mm + ne*em + nmd*mmd + nh*hm + nn*nmm
    p(); rule("═"); ctr("MARKS SUMMARY"); rule()
    p(f"  A(MCQs):{nm*mm}  B(Short):{ne*em}  C(Desc):{nmd*mmd}  "
      f"D(Hard):{nh*hm}  E(Num):{nn*nmm}  |  TOTAL:{actual}")
    if actual != marks:
        p(f"  ⚠ Actual={actual} vs Set={marks}. Adjust marks per question if needed.")
    rule("═")

    if show_k:
        p(); p(); rule("═")
        ctr("★  ANSWER KEY  (Teacher Copy)  ★")
        rule("═"); p()
        if mcqs and ak.get("mcq_answers"):
            p("MCQ ANSWERS:"); rule()
            for i, m in enumerate(mcqs):
                let = chr(65 + m.get("correct_index", 0))
                ans = ak["mcq_answers"][i] if i < len(ak["mcq_answers"]) else m.get("correct_text", "")
                p(f"  Q{i+1}. ({let}) {ans}")
            p()
        hints_data = [
            ("SECTION B HINTS:", "easy_hints", ne),
            ("SECTION C HINTS:", "medium_hints", nmd),
            ("SECTION D HINTS:", "hard_hints", nh),
        ]
        offset = nm
        for title_h, key, count in hints_data:
            hints = ak.get(key, [])
            if hints:
                p(title_h); rule()
                for i in range(count):
                    h = hints[i] if i < len(hints) else "See document."
                    p(f"  Q{offset+i+1}. {h}")
                offset += count
                p()
        if nums:
            p("NUMERICAL HINTS:"); rule()
            ns = nm + ne + nmd + nh
            for i, n in enumerate(nums):
                p(f"  Q{ns+i+1}. {n.get('solution_hint','Apply relevant formula.')}")
            p()
        rule("═")

    return "\n".join(L)

# ═══════════════════════════════════════════════════════════════════════
#  BUILT-IN QUESTION BANK
# ═══════════════════════════════════════════════════════════════════════
BANK = {
    "Physics": {
        "easy":   ["Define velocity and state its SI unit.",
                   "What is Newton's First Law of Motion? Give one example.",
                   "Define scalar and vector quantities with two examples each.",
                   "State Ohm's Law. Write its mathematical form.",
                   "What is kinetic energy? Write its formula.",
                   "Define electric current and state its SI unit.",
                   "What is the speed of light in vacuum?",
                   "Define work done by a force. State its SI unit."],
        "medium": ["Explain the difference between speed and velocity. Why is velocity a vector?",
                   "Describe the three equations of uniformly accelerated motion and derive one.",
                   "Explain refraction of light. State Snell's Law with a real example.",
                   "Explain series vs parallel circuits with household wiring examples.",
                   "Describe how sound waves propagate through different media."],
        "hard":   ["Analyze why a satellite doesn't fall to Earth despite free-fall. Derive orbital velocity.",
                   "Ohm's Law breaks down for semiconductors — critically analyze its limitations.",
                   "Explain wave-particle duality and de Broglie's hypothesis with experimental evidence."],
        "mcq":    [("SI unit of power:", ["Joule","Watt","Newton","Pascal"], 1),
                   ("F = ma is stated by:", ["Newton's 1st","Newton's 2nd","Newton's 3rd","Hooke's"], 1),
                   ("Speed of light (vacuum):", ["3×10⁶ m/s","3×10⁸ m/s","3×10¹⁰","3×10⁴"], 1),
                   ("Velocity is a:", ["Scalar","Vector","Constant","Force"], 1),
                   ("KE formula:", ["mgh","½mv²","mv","Fs"], 1),
                   ("Unit of resistance:", ["Volt","Ampere","Ohm","Farad"], 2),
                   ("Concave mirror headlights:", ["Plane","Concave","Convex","Parabolic"], 1),],
        "num":    [{"q":"Car from rest, a={a}m/s² for t={t}s. Find velocity and distance.",
                    "vars":{"a":(2,8),"t":(3,10)},"formula":"v=u+at; s=½at²",
                    "sol":lambda v:f"v={v['a']*v['t']}m/s, s={0.5*v['a']*v['t']**2:.1f}m"},
                   {"q":"Mass {m}kg at {vel}m/s. Find KE.",
                    "vars":{"m":(2,20),"vel":(3,15)},"formula":"KE=½mv²",
                    "sol":lambda v:f"KE={0.5*v['m']*v['vel']**2:.1f}J"}]
    },
    "Mathematics": {
        "easy":   ["Define a rational number with two examples.",
                   "State the Pythagorean theorem.",
                   "Write formulae for area and circumference of a circle.",
                   "Find HCF of 24 and 36.",
                   "What is a prime number? List all primes up to 20.",
                   "Define absolute value. Find |−7| and |3−9|."],
        "medium": ["Solve using substitution: 2x+y=10 and x−y=2.",
                   "Derive the quadratic formula by completing the square.",
                   "Explain the Chain Rule of differentiation with two examples.",
                   "State laws of logarithms. Simplify: log100 + log½ − log5.",
                   "Explain classical probability with examples."],
        "hard":   ["Prove √2 is irrational by contradiction. Generalize when √n is irrational.",
                   "State and prove the Fundamental Theorem of Calculus.",
                   "Analyze how 'a' in ax²+bx+c affects discriminant, vertex, and roots."],
        "mcq":    [("π to 2 decimal places:", ["3.14","3.41","3.12","3.16"], 0),
                   ("Derivative of x²:", ["x","2x","2","x³/3"], 1),
                   ("Angles in a triangle:", ["90°","180°","270°","360°"], 1),
                   ("Area of circle:", ["πr","2πr","πr²","2πr²"], 2),
                   ("log(1) =", ["0","1","-1","undefined"], 0),
                   ("∫x dx =", ["x²","x²/2+C","2x+C","x+C"], 1),
                   ("Cos 90° =", ["1","0","-1","√2/2"], 1),],
        "num":    [{"q":"Roots of x²−{b}x+{c}=0?","vars":{"b":(5,12),"c":(4,20)},
                    "formula":"x=(b±√(b²−4c))/2",
                    "sol":lambda v:(lambda D:f"D={D}; x=({v['b']}±√{D})/2")(v['b']**2-4*v['c'])},
                   {"q":"CI on ₹{P} at {r}% for {t} years.","vars":{"P":(1000,10000),"r":(5,15),"t":(2,5)},
                    "formula":"A=P(1+r/100)^n; CI=A−P",
                    "sol":lambda v:(lambda A:f"A=₹{A:.2f}; CI=₹{A-v['P']:.2f}")(v['P']*(1+v['r']/100)**v['t'])}]
    },
    "Chemistry": {
        "easy":   ["Define an atom and a molecule with examples.",
                   "State the Law of Conservation of Mass.",
                   "What is pH of pure water? What does pH<7 mean?",
                   "Define oxidation and reduction in terms of electrons.",
                   "State Avogadro's number and its value.",
                   "What is the atomic number of Carbon?"],
        "medium": ["Compare ionic and covalent bonding with two examples each.",
                   "Describe electrolysis of water with electrode reactions.",
                   "Identify oxidizing/reducing agents: Zn + CuSO₄ → ZnSO₄ + Cu.",
                   "Explain collision theory of reaction rates.",
                   "Describe allotropes of carbon and compare their structures."],
        "hard":   ["Noble gases form compounds despite the octet rule — critically analyze.",
                   "Analyze ΔG=ΔH−TΔS to explain why some reactions are spontaneous only at high T.",
                   "Analyze Le Chatelier's Principle for N₂+3H₂⇌2NH₃."],
        "mcq":    [("Atomic number of Carbon:", ["4","6","8","12"], 1),
                   ("pH of neutral:", ["0","7","14","1"], 1),
                   ("NaCl bond type:", ["Covalent","Ionic","Metallic","Hydrogen"], 1),
                   ("Lightest element:", ["Helium","Hydrogen","Lithium","Carbon"], 1),
                   ("Oxidation =", ["Gain e⁻","Loss e⁻","Gain H⁺","Loss H⁺"], 1),
                   ("Formula of water:", ["H₂O₂","HO","H₂O","H₃O⁺"], 2),
                   ("Zn+HCl gas:", ["O₂","Cl₂","H₂","CO₂"], 2),],
        "num":    [{"q":"Moles in {mass}g of H₂O (M=18g/mol)?","vars":{"mass":(18,180)},
                    "formula":"moles=mass/M",
                    "sol":lambda v:f"moles={v['mass']}/18={v['mass']/18:.2f}mol"},
                   {"q":"Mass of {m} moles of CO₂ (M=44)?","vars":{"m":(1,10)},
                    "formula":"mass=moles×M",
                    "sol":lambda v:f"mass={v['m']}×44={v['m']*44}g"}]
    },
    "Computer Science": {
        "easy":   ["What is an algorithm? List its key properties.",
                   "Difference between compiler and interpreter?",
                   "Define RAM and ROM.",
                   "Convert decimal 10 to binary.",
                   "What is a function/subroutine?",
                   "Define a database. Give one DBMS example."],
        "medium": ["Explain OOP's four pillars with examples.",
                   "Trace bubble sort and selection sort on [5,3,8,1,2].",
                   "Explain how TCP/IP works layer by layer.",
                   "Explain recursion with factorial. What are the risks?",
                   "Describe stack and queue with real applications."],
        "hard":   ["Analyze QuickSort complexity — best/average/worst. Why preferred?",
                   "Analyze Coffman's deadlock conditions with three prevention strategies.",
                   "How does garbage collection work? Compare GC strategies."],
        "mcq":    [("CPU =", ["Central Processing Unit","Computer Processing Unit","Core Processing Unit","Control Program Unit"], 0),
                   ("Binary of 10:", ["1100","1010","1001","1110"], 1),
                   ("Not OOP:", ["Encapsulation","Polymorphism","Compilation","Inheritance"], 2),
                   ("RAM is:", ["Non-volatile","Volatile","Permanent","Read-only"], 1),
                   ("O(n log n) sort:", ["Bubble","Selection","QuickSort","Insertion"], 2),
                   ("Stack uses:", ["FIFO","LIFO","Random","Priority"], 1),
                   ("IPv4 bits:", ["16","32","64","128"], 1),],
        "num":    [{"q":"Convert {n} decimal to binary, octal, hex.","vars":{"n":(10,255)},
                    "formula":"Divide by 2/8/16 repeatedly",
                    "sol":lambda v:f"Bin:{bin(v['n'])[2:]} Oct:{oct(v['n'])[2:]} Hex:{hex(v['n'])[2:].upper()}"}]
    }
}

def build_builtin_paper(config):
    subj  = config.get("subject","Physics")
    bank  = BANK.get(subj, BANK["Physics"])
    ne, nm2, nh = config.get("easy",3), config.get("medium",3), config.get("hard",2)
    nmcq = config.get("mcq_count",5)
    nnum = config.get("numerical_count",2)
    marks = config.get("total_marks",100)

    def pick(lst, n): return random.sample(lst, min(n, len(lst)))
    mcq_raw = pick(bank["mcq"], nmcq)
    mcqs = []
    for stem, opts, ci in mcq_raw:
        correct = opts[ci]; sh = opts[:]; random.shuffle(sh)
        mcqs.append({"question":stem,"options":sh,
                     "correct_index":sh.index(correct),"correct_text":correct})

    num_data = bank.get("num",[])
    nums = []
    for item in pick(num_data, nnum):
        vals = {k:random.randint(*r) for k,r in item["vars"].items()}
        try: sol = item["sol"](vals)
        except: sol = "See formula sheet."
        nums.append({"question":item["q"].format(**vals),"formula":item["formula"],"solution_hint":sol})

    ai_data = {
        "subject_detected": subj,
        "topic_summary": f"Built-in {subj} question bank.",
        "mcqs": mcqs,
        "easy_questions": pick(bank["easy"], ne),
        "medium_questions": pick(bank["medium"], nm2),
        "hard_questions": pick(bank["hard"], nh),
        "numerical_problems": nums,
        "answer_key": {
            "mcq_answers": [m["correct_text"] for m in mcqs],
            "easy_hints": ["Define clearly with examples."]*ne,
            "medium_hints": ["Explain with 3–5 key points."]*nm2,
            "hard_hints": ["Analytical — use evidence."]*nh,
        }
    }
    return format_exam_paper(ai_data, config, source_label="Built-in Question Bank")

# ═══════════════════════════════════════════════════════════════════════
#  MAIN APP
# ═══════════════════════════════════════════════════════════════════════
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("EduBudget Pro v3")
        self.geometry("1320x820")
        self.minsize(1050, 680)
        self.configure(bg=BG)
        self.data = load_data()
        self._pages = {}
        self._cur   = None
        self._build()
        self._nav("dashboard")

    def _build(self):
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)
        self._sidebar()
        self.content = tk.Frame(self, bg=BG)
        self.content.grid(row=0, column=1, sticky="nsew")
        self.content.grid_rowconfigure(0, weight=1)
        self.content.grid_columnconfigure(0, weight=1)

    def _sidebar(self):
        sb = tk.Frame(self, bg=BG, width=230)
        sb.grid(row=0, column=0, sticky="nsew")
        sb.grid_propagate(False)

        tk.Label(sb, text="⚡ EduBudget", font=("Segoe UI",16,"bold"),
                 fg=ACCENT, bg=BG).pack(anchor="w", padx=18, pady=(24,0))
        tk.Label(sb, text="Pro  v3", font=("Segoe UI",14,"bold"),
                 fg=AI_CLR, bg=BG).pack(anchor="w", padx=18)
        tk.Label(sb, text="Student Productivity System", font=FS,
                 fg=DIM, bg=BG).pack(anchor="w", padx=18)
        tk.Frame(sb, bg=BORDER, height=1).pack(fill="x", padx=14, pady=14)

        self._nav_btns = {}
        for key, label in [("dashboard","🏠  Dashboard"),
                             ("exam",     "📄  Exam Generator"),
                             ("budget",   "💰  Budget Manager"),
                             ("history",  "📊  History")]:
            b = tk.Button(sb, text=label, font=FB, anchor="w",
                          bg=BG, fg=TEXT, relief="flat", bd=0,
                          padx=20, pady=12, cursor="hand2",
                          activebackground=SEL, activeforeground="white",
                          command=lambda k=key: self._nav(k))
            b.pack(fill="x", padx=8, pady=2)
            self._nav_btns[key] = b

        tk.Label(sb, text="v3.0", font=FS, fg=DIM, bg=BG).pack(side="bottom", pady=12)

    def _nav(self, key):
        if self._cur == key:
            return
        self._cur = key
        for k, b in self._nav_btns.items():
            b.config(bg=SEL if k==key else BG, fg="white" if k==key else TEXT)
        if key not in self._pages:
            cls = {"dashboard":DashboardPage,"exam":ExamPage,
                   "budget":BudgetPage,"history":HistoryPage}[key]
            p = cls(self.content, self)
            p.grid(row=0, column=0, sticky="nsew")
            self._pages[key] = p
        else:
            self._pages[key].refresh()
        self._pages[key].tkraise()

    def save(self):
        save_data(self.data)

    def go(self, key):
        self._nav(key)


# ═══════════════════════════════════════════════════════════════════════
#  DASHBOARD
# ═══════════════════════════════════════════════════════════════════════
class DashboardPage(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._build()

    def _build(self):
        # Header
        h = tk.Frame(self, bg=BG)
        h.pack(fill="x", padx=32, pady=(28,12))
        tk.Label(h, text="Dashboard", font=FT, fg=BRIGHT, bg=BG).pack(anchor="w")
        tk.Label(h, text=datetime.datetime.now().strftime("%A, %d %B %Y"),
                 font=F, fg=DIM, bg=BG).pack(anchor="w")

        # Stat cards
        row = tk.Frame(self, bg=BG)
        row.pack(fill="x", padx=28, pady=8)
        self._stats = {}
        for key, icon, label, clr in [
            ("exams",  "📄", "Exams Generated",    ACCENT),
            ("exp",    "💸", "Month Expenses",      ORANGE),
            ("top",    "🏆", "Top Spend Category",  PURPLE),
            ("status", "💡", "Budget Status",       GREEN),
        ]:
            f = tk.Frame(row, bg=CARD, padx=18, pady=16)
            f.pack(side="left", expand=True, fill="both", padx=8)
            tk.Label(f, text=icon, font=("Segoe UI",22), bg=CARD, fg=clr).pack(anchor="w")
            v = tk.Label(f, text="—", font=("Segoe UI",17,"bold"), fg=clr, bg=CARD)
            v.pack(anchor="w")
            tk.Label(f, text=label, font=FS, fg=DIM, bg=CARD).pack(anchor="w")
            self._stats[key] = v

        # Quick actions
        qa = tk.Frame(self, bg=BG)
        qa.pack(fill="x", padx=28, pady=(14,6))
        tk.Label(qa, text="Quick Actions", font=FM, fg=ACCENT, bg=BG).pack(anchor="w", pady=(0,8))
        br = tk.Frame(qa, bg=BG)
        br.pack(anchor="w")
        for label, page, color in [
            ("📄 Generate Exam", "exam",    ACCENT),
            ("💰 Budget",        "budget",  ORANGE),
            ("📊 History",       "history", PURPLE),
        ]:
            btn(br, label, cmd=lambda p=page: self.app.go(p),
                bg=color, padx=16, pady=10).pack(side="left", padx=(0,10))

        # Info panels
        info = tk.Frame(self, bg=BG)
        info.pack(fill="both", expand=True, padx=28, pady=10)

        left = tk.Frame(info, bg=CARD, padx=18, pady=14)
        left.pack(side="left", fill="both", expand=True, padx=(0,8))
        tk.Label(left, text="💰 Budget Snapshot", font=FM, fg=ACCENT, bg=CARD).pack(anchor="w")
        tk.Frame(left, bg=BORDER, height=1).pack(fill="x", pady=6)
        self._snap = tk.Text(left, bg=CARD, fg=TEXT, font=F, relief="flat",
                              state="disabled", height=10, wrap="word")
        self._snap.pack(fill="both", expand=True)

        right = tk.Frame(info, bg=CARD, padx=18, pady=14, width=280)
        right.pack(side="right", fill="both")
        right.pack_propagate(False)
        tk.Label(right, text="📋 What's New in v3", font=FM, fg=AI_CLR, bg=CARD).pack(anchor="w")
        tk.Frame(right, bg=BORDER, height=1).pack(fill="x", pady=6)
        guide = ("✅ FIXED: Budget section — clean & stable\n\n"
                 "✨ NEW: Smart Expense Split\n"
                 "  Enter income + savings goal\n"
                 "  → Tick categories you need\n"
                 "  → Auto-split using 50/30/20 rule\n"
                 "  → Edit any amount freely\n\n"
                 "📄 PDF/Word → Questions\n"
                 "  • Offline mode — no API key needed!\n"
                 "  • Extracts text & auto-generates Qs\n"
                 "  • Optional: Claude AI for smarter Qs\n\n"
                 "📄 Built-in question bank still works\n"
                 "   for Physics, Math, Chemistry, CS")
        t = tk.Text(right, bg=CARD, fg=TEXT, font=FS, relief="flat",
                    state="disabled", wrap="word")
        t.pack(fill="both", expand=True)
        t.config(state="normal"); t.insert("end", guide); t.config(state="disabled")
        self.refresh()

    def refresh(self):
        d = self.app.data
        a = calc_budget(d.get("budget", {}))
        self._stats["exams"].config(text=str(d.get("exams_generated", 0)))
        self._stats["exp"].config(text=f"₹{a['total_exp']:,.0f}")
        self._stats["top"].config(text=a["top_cat"])
        self._stats["status"].config(
            text="⚠ Overspending" if a["over"] else "✅ On Track",
            fg=RED if a["over"] else GREEN)
        self._snap.config(state="normal")
        self._snap.delete("1.0","end")
        self._snap.insert("end",
            f"Income:        ₹{a['income']:>10,.2f}\n"
            f"Savings Goal:  ₹{a['savings_goal']:>10,.2f}\n"
            f"Total Spent:   ₹{a['total_exp']:>10,.2f}\n"
            f"Remaining:     ₹{a['remaining']:>10,.2f}\n\n"
            "── Category Totals ─────────────────\n")
        for cat, amt in sorted(a["cat_totals"].items(), key=lambda x: -x[1]):
            if amt > 0:
                self._snap.insert("end", f"  {cat:<14} ₹{amt:>8,.2f}\n")
        self._snap.config(state="disabled")


# ═══════════════════════════════════════════════════════════════════════
#  EXAM GENERATOR
# ═══════════════════════════════════════════════════════════════════════
class ExamPage(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._paper = ""
        self._doc_text = None
        self._busy = False
        self._build()

    def _build(self):
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)

        # Left panel
        left = tk.Frame(self, bg=CARD, width=350)
        left.grid(row=0, column=0, sticky="nsew")
        left.grid_propagate(False)

        # Tab strip
        tabs = tk.Frame(left, bg=INPUT)
        tabs.pack(fill="x")
        self._tab = tk.StringVar(value="builtin")
        self._tb = {}
        for key, label in [("builtin","📚 Built-in"), ("doc","📄 From Document")]:
            b = tk.Button(tabs, text=label, font=FB, bg=CARD if key=="builtin" else INPUT,
                          fg=ACCENT if key=="builtin" else DIM,
                          relief="flat", bd=0, padx=12, pady=10, cursor="hand2",
                          command=lambda k=key: self._switch(k))
            b.pack(side="left", expand=True, fill="x")
            self._tb[key] = b

        self._f_builtin = tk.Frame(left, bg=CARD)
        self._f_builtin.pack(fill="both", expand=True)
        self._f_doc = tk.Frame(left, bg=CARD)

        self._build_builtin(self._f_builtin)
        self._build_doc(self._f_doc)

        # Right: preview
        right = tk.Frame(self, bg=BG)
        right.grid(row=0, column=1, sticky="nsew", padx=(3,0))
        right.grid_rowconfigure(1, weight=1)
        right.grid_columnconfigure(0, weight=1)

        hdr = tk.Frame(right, bg=BG)
        hdr.grid(row=0, column=0, sticky="ew", padx=18, pady=(18,6))
        tk.Label(hdr, text="Exam Paper Preview", font=FH, fg=BRIGHT, bg=BG).pack(side="left")
        self._status_lbl = tk.Label(hdr, text="No paper generated.", font=FS, fg=DIM, bg=BG)
        self._status_lbl.pack(side="right")

        pf = tk.Frame(right, bg=CARD)
        pf.grid(row=1, column=0, sticky="nsew", padx=18, pady=(0,18))
        pf.grid_rowconfigure(0, weight=1)
        pf.grid_columnconfigure(0, weight=1)
        self._preview = scrolledtext.ScrolledText(
            pf, bg=BG, fg=TEXT, font=("Courier New",9),
            insertbackground=TEXT, relief="flat", wrap="word", padx=10, pady=10)
        self._preview.grid(row=0, column=0, sticky="nsew")
        self._set_preview("Configure on the left and click Generate.")

    def _switch(self, tab):
        self._tab.set(tab)
        if tab == "builtin":
            self._f_builtin.pack(fill="both", expand=True)
            self._f_doc.pack_forget()
            self._tb["builtin"].config(bg=CARD, fg=ACCENT)
            self._tb["doc"].config(bg=INPUT, fg=DIM)
        else:
            self._f_doc.pack(fill="both", expand=True)
            self._f_builtin.pack_forget()
            self._tb["doc"].config(bg=CARD, fg=ORANGE)
            self._tb["builtin"].config(bg=INPUT, fg=DIM)

    def _cfg_fields(self, parent):
        """Common paper config fields. Returns dict of vars."""
        v = {k: tk.StringVar() for k in ["school","title","teacher","dur"]}
        v["school"].set("My School")
        v["title"].set("Examination")
        v["teacher"].set("")
        v["dur"].set("2 Hours")
        v["marks"]  = tk.IntVar(value=100)
        v["easy"]   = tk.IntVar(value=4)
        v["medium"] = tk.IntVar(value=3)
        v["hard"]   = tk.IntVar(value=2)
        v["mcq"]    = tk.IntVar(value=8)
        v["num"]    = tk.IntVar(value=2)
        v["key"]    = tk.BooleanVar(value=True)

        rows = [("School:", "school", "s"), ("Title:", "title", "s"),
                ("Teacher:", "teacher", "s"), ("Duration:", "dur", "s"),
                ("Total Marks:", "marks", "n"), ("Easy Qs:", "easy", "n"),
                ("Medium Qs:", "medium", "n"), ("Hard Qs:", "hard", "n"),
                ("MCQs:", "mcq", "n"), ("Numericals:", "num", "n")]

        for ltext, key, typ in rows:
            f = tk.Frame(parent, bg=CARD)
            f.pack(fill="x", padx=18, pady=2)
            tk.Label(f, text=ltext, font=F, fg=DIM, bg=CARD,
                     width=14, anchor="w").pack(side="left")
            if typ == "s":
                entry(f, v[key], 16).pack(side="left")
            else:
                spin(f, v[key], 0, 300, 8).pack(side="left")

        ck = tk.Checkbutton(parent, text=" Include Answer Key", variable=v["key"],
                             font=F, fg=TEXT, bg=CARD, activebackground=CARD,
                             selectcolor=INPUT)
        ck.pack(anchor="w", padx=18, pady=4)
        return v

    def _build_builtin(self, parent):
        inner = scrollable(parent)
        tk.Label(inner, text="📚 Built-in Question Bank", font=FH,
                 fg=ACCENT, bg=CARD).pack(anchor="w", padx=18, pady=(18,4))
        sep(inner, pady=4)

        # Subject
        tk.Label(inner, text="Subject", font=FM, fg=BRIGHT, bg=CARD).pack(anchor="w", padx=18, pady=(10,4))
        sf = tk.Frame(inner, bg=CARD)
        sf.pack(fill="x", padx=18)
        self._v_subj = tk.StringVar(value="Physics")
        ttk.Combobox(sf, textvariable=self._v_subj,
                     values=list(BANK.keys()), state="readonly",
                     width=18, font=F).pack(side="left")

        tk.Label(inner, text="Paper Settings", font=FM, fg=BRIGHT, bg=CARD).pack(anchor="w", padx=18, pady=(12,4))
        self._bv = self._cfg_fields(inner)

        sep(inner, pady=8)
        btn(inner, "⚡ Generate Paper", cmd=self._gen_builtin,
            bg=ACCENT, padx=18, pady=10).pack(fill="x", padx=18, pady=4)
        btn(inner, "💾 Export TXT", cmd=self._export,
            bg=GREEN, padx=18, pady=8).pack(fill="x", padx=18, pady=4)
        btn(inner, "🗑 Clear", cmd=self._clear,
            bg=INPUT, fg=DIM, padx=18, pady=6).pack(fill="x", padx=18, pady=4)
        tk.Label(inner, bg=CARD).pack(pady=16)

    def _build_doc(self, parent):
        inner = scrollable(parent)
        tk.Label(inner, text="📄 Questions from Document", font=FH,
                 fg=ORANGE, bg=CARD).pack(anchor="w", padx=18, pady=(18,4))
        tk.Label(inner, text="Upload PDF or Word → auto-generate questions\n(works offline! no API key needed)",
                 font=FS, fg=DIM, bg=CARD, justify="left").pack(anchor="w", padx=18)
        sep(inner, pady=6)

        # Step 1: Upload
        tk.Label(inner, text="Step 1 — Upload Document", font=FM,
                 fg=BRIGHT, bg=CARD).pack(anchor="w", padx=18, pady=(4,4))
        btn(inner, "📁 Browse PDF / Word (.docx)",
            cmd=self._browse, bg=PURPLE, padx=14, pady=8).pack(anchor="w", padx=18, pady=4)
        self._file_lbl  = tk.Label(inner, text="No file selected", font=FS, fg=DIM, bg=CARD, wraplength=290)
        self._file_lbl.pack(anchor="w", padx=18)
        self._file_info = tk.Label(inner, text="", font=FS, fg=GREEN, bg=CARD, wraplength=290, justify="left")
        self._file_info.pack(anchor="w", padx=18)

        sep(inner, pady=6)

        # Step 2: Mode
        tk.Label(inner, text="Step 2 — Generation Mode", font=FM,
                 fg=BRIGHT, bg=CARD).pack(anchor="w", padx=18, pady=(4,4))
        self._v_mode = tk.StringVar(value="offline")
        mf = tk.Frame(inner, bg=CARD)
        mf.pack(fill="x", padx=18, pady=4)
        tk.Radiobutton(mf, text="🔵 Offline (free, no key needed)",
                       variable=self._v_mode, value="offline",
                       font=F, fg=GREEN, bg=CARD, activebackground=CARD,
                       selectcolor=INPUT, command=self._mode_changed).pack(anchor="w")
        tk.Radiobutton(mf, text="🤖 Claude AI (smarter, needs API key)",
                       variable=self._v_mode, value="ai",
                       font=F, fg=AI_CLR, bg=CARD, activebackground=CARD,
                       selectcolor=INPUT, command=self._mode_changed).pack(anchor="w")

        # API key (hidden by default)
        self._api_frame = tk.Frame(inner, bg=CARD)
        ak_row = tk.Frame(self._api_frame, bg=CARD)
        ak_row.pack(fill="x", padx=18, pady=4)
        tk.Label(ak_row, text="API Key:", font=F, fg=DIM, bg=CARD, width=10, anchor="w").pack(side="left")
        self._v_api = tk.StringVar(value=self.app.data.get("api_key",""))
        self._api_ent = entry(ak_row, self._v_api, 20, show="•")
        self._api_ent.pack(side="left", padx=(0,4))
        tk.Button(ak_row, text="👁", font=FS, bg=INPUT, fg=DIM, relief="flat",
                  cursor="hand2",
                  command=lambda: self._api_ent.config(
                      show="" if self._api_ent["show"]=="•" else "•")
                  ).pack(side="left")
        tk.Label(self._api_frame, text="  Get free key: console.anthropic.com",
                 font=FS, fg=DIM, bg=CARD).pack(anchor="w", padx=18)
        btn(self._api_frame, "💾 Save Key", cmd=self._save_key,
            bg=INPUT, fg=DIM, padx=10, pady=5).pack(anchor="w", padx=18, pady=4)

        sep(inner, pady=6)

        # Step 3: Paper config
        tk.Label(inner, text="Step 3 — Paper Settings", font=FM,
                 fg=BRIGHT, bg=CARD).pack(anchor="w", padx=18, pady=(4,4))
        self._dv = self._cfg_fields(inner)
        self._dv_subj_hint = tk.StringVar(value="Auto-detect")
        f2 = tk.Frame(inner, bg=CARD)
        f2.pack(fill="x", padx=18, pady=2)
        tk.Label(f2, text="Subject Hint:", font=F, fg=DIM, bg=CARD,
                 width=14, anchor="w").pack(side="left")
        entry(f2, self._dv_subj_hint, 16).pack(side="left")

        sep(inner, pady=8)
        self._gen_btn = btn(inner, "🚀 Generate Questions",
                             cmd=self._gen_doc, bg=ORANGE, fg="white",
                             padx=18, pady=10)
        self._gen_btn.pack(fill="x", padx=18, pady=4)
        self._prog_lbl = tk.Label(inner, text="", font=FB, fg=ORANGE, bg=CARD, wraplength=290)
        self._prog_lbl.pack(anchor="w", padx=18, pady=4)
        btn(inner, "💾 Export TXT", cmd=self._export,
            bg=GREEN, padx=18, pady=8).pack(fill="x", padx=18, pady=4)
        btn(inner, "🗑 Clear", cmd=self._clear,
            bg=INPUT, fg=DIM, padx=18, pady=6).pack(fill="x", padx=18, pady=4)
        tk.Label(inner, bg=CARD).pack(pady=16)

    def _mode_changed(self):
        if self._v_mode.get() == "ai":
            self._api_frame.pack(fill="x", after=None)
            # find the sep after mode frame and insert before it
            self._api_frame.pack(fill="x", padx=0, pady=0)
        else:
            self._api_frame.pack_forget()

    def _browse(self):
        path = filedialog.askopenfilename(
            title="Select Study Material",
            filetypes=[("PDF/Word","*.pdf *.docx"),
                       ("PDF","*.pdf"),("Word","*.docx"),("All","*.*")])
        if not path: return
        self._file_lbl.config(text=os.path.basename(path), fg=TEXT)
        self._file_info.config(text="⏳ Reading...", fg=YELLOW)
        self.update()
        try:
            text = extract_text(path)
            self._doc_text = text
            words = len(text.split())
            self._file_info.config(
                text=f"✅ {words:,} words extracted. Ready!", fg=GREEN)
        except Exception as e:
            self._doc_text = None
            self._file_info.config(text=f"❌ {e}", fg=RED)

    def _save_key(self):
        k = self._v_api.get().strip()
        if not k:
            messagebox.showwarning("Empty","Enter API key first.")
            return
        self.app.data["api_key"] = k
        self.app.save()
        messagebox.showinfo("Saved","API key saved.")

    def _get_config(self, vars_dict, subject=""):
        return {
            "subject": subject or vars_dict.get("subj_hint", tk.StringVar()).get() or "Auto",
            "title":   vars_dict["title"].get() or "Examination",
            "school_name": vars_dict["school"].get(),
            "teacher_name": vars_dict["teacher"].get(),
            "duration": vars_dict["dur"].get(),
            "total_marks": vars_dict["marks"].get(),
            "easy":    vars_dict["easy"].get(),
            "medium":  vars_dict["medium"].get(),
            "hard":    vars_dict["hard"].get(),
            "mcq_count": vars_dict["mcq"].get(),
            "numerical_count": vars_dict["num"].get(),
            "show_answer_key": vars_dict["key"].get(),
        }

    def _gen_builtin(self):
        cfg = self._get_config(self._bv, subject=self._v_subj.get())
        if cfg["easy"]+cfg["medium"]+cfg["hard"]+cfg["mcq_count"]+cfg["numerical_count"] == 0:
            messagebox.showwarning("No questions","Set at least one question count > 0.")
            return
        paper = build_builtin_paper(cfg)
        self._paper = paper
        self._set_preview(paper)
        self._status_lbl.config(text=f"✅ {datetime.datetime.now().strftime('%H:%M:%S')}", fg=GREEN)
        self.app.data["exams_generated"] = self.app.data.get("exams_generated",0) + 1
        self.app.save()

    def _gen_doc(self):
        if self._busy:
            messagebox.showinfo("Busy","Generation in progress, please wait.")
            return
        if not self._doc_text:
            messagebox.showerror("No Document","Upload a PDF or Word file first.")
            return
        cfg = self._get_config(self._dv)
        cfg["subject"] = self._dv_subj_hint.get() or "Auto-detect"
        if cfg["easy"]+cfg["medium"]+cfg["hard"]+cfg["mcq_count"] == 0:
            messagebox.showwarning("No questions","Set at least one question count > 0.")
            return

        mode = self._v_mode.get()
        if mode == "ai":
            api_key = self._v_api.get().strip()
            if not api_key:
                messagebox.showerror("No API Key",
                    "Enter your Anthropic API key in Step 2.\n"
                    "Or switch to Offline mode — it's free!")
                return
            self._busy = True
            self._gen_btn.config(state="disabled", text="⏳ Generating...")
            self._prog_lbl.config(text="🤖 Sending to Claude AI...", fg=AI_CLR)
            self._set_preview("⏳ Claude AI is reading your document...\n\nPlease wait 15–40 seconds.")

            def run():
                try:
                    def prog(m): self.after(0, lambda: self._prog_lbl.config(text=f"🤖 {m}"))
                    ai_data = ai_generate_questions(self._doc_text, cfg, api_key, on_progress=prog)
                    paper   = format_exam_paper(ai_data, cfg, "AI — Claude")
                    def done():
                        self._paper = paper
                        self._set_preview(paper)
                        self._status_lbl.config(text=f"✅ AI Done {datetime.datetime.now().strftime('%H:%M:%S')}", fg=GREEN)
                        self._prog_lbl.config(text="✅ Generated successfully!", fg=GREEN)
                        self._busy = False
                        self._gen_btn.config(state="normal", text="🚀 Generate Questions")
                        self.app.data["exams_generated"] = self.app.data.get("exams_generated",0) + 1
                        self.app.save()
                    self.after(0, done)
                except Exception as e:
                    def err():
                        self._busy = False
                        self._gen_btn.config(state="normal", text="🚀 Generate Questions")
                        self._prog_lbl.config(text=f"❌ {str(e)[:100]}", fg=RED)
                        self._set_preview(f"❌ AI Error: {e}\n\nTry Offline mode instead — it's free!")
                    self.after(0, err)
            threading.Thread(target=run, daemon=True).start()
        else:
            # Offline generation
            self._prog_lbl.config(text="⚙ Generating offline...", fg=GREEN)
            self.update()
            try:
                ai_data = offline_generate_questions(self._doc_text, cfg)
                paper   = format_exam_paper(ai_data, cfg, "Offline — Keyword Extraction")
                self._paper = paper
                self._set_preview(paper)
                self._status_lbl.config(text=f"✅ Offline {datetime.datetime.now().strftime('%H:%M:%S')}", fg=GREEN)
                self._prog_lbl.config(text="✅ Done (offline mode)!", fg=GREEN)
                self.app.data["exams_generated"] = self.app.data.get("exams_generated",0) + 1
                self.app.save()
            except Exception as e:
                self._prog_lbl.config(text=f"❌ {e}", fg=RED)
                messagebox.showerror("Error", str(e))

    def _set_preview(self, text):
        self._preview.config(state="normal")
        self._preview.delete("1.0","end")
        self._preview.insert("end", text)
        self._preview.config(state="disabled")

    def _export(self):
        if not self._paper:
            messagebox.showinfo("No Paper","Generate a paper first.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text","*.txt"),("All","*.*")],
            initialfile=f"Exam_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.txt")
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self._paper)
            messagebox.showinfo("Saved", f"Saved to:\n{path}")

    def _clear(self):
        self._paper = ""
        self._set_preview("Paper cleared. Generate a new one.")
        self._status_lbl.config(text="Cleared.", fg=DIM)
        self._prog_lbl.config(text="")

    def refresh(self):
        self._v_api.set(self.app.data.get("api_key",""))


# ═══════════════════════════════════════════════════════════════════════
#  BUDGET PAGE — Clean rewrite with smart split flow
# ═══════════════════════════════════════════════════════════════════════
class BudgetPage(tk.Frame):
    """
    FLOW:
      Step 1 → Enter income + savings goal + month
      Step 2 → Tick categories you spend on
      Step 3 → Click "Suggest Split" → amounts auto-filled (editable)
      Step 4 → Add custom expenses if needed
      Step 5 → Calculate → see analysis

    FIX: All Entry widgets stored as direct refs.
         _sync() reads from widget.get() NOT StringVars
         to avoid canvas/StringVar binding bug on Windows.
    """
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._cat_vars    = {}   # category -> BooleanVar (checkbox)
        self._amt_entries = {}   # category -> Entry widget (direct ref)
        self._custom_list = []   # [{name, amount, category}]
        # Direct Entry refs for top fields (avoids StringVar canvas bug)
        self._ent_month   = None
        self._ent_income  = None
        self._ent_savings = None
        self._ent_cname   = None
        self._ent_camt    = None
        self._v_ccat      = tk.StringVar(value="Food")
        self._build()
        self._load()

    def _make_entry(self, parent, default="", width=14, row=None, col=None):
        """Create Entry widget, store NO StringVar — read .get() directly."""
        e = tk.Entry(parent, width=width, bg=INPUT, fg=BRIGHT,
                     insertbackground=BRIGHT, relief="flat", font=F, bd=4)
        e.insert(0, default)
        if row is not None:
            e.grid(row=row, column=col, padx=(4,0), sticky="w", pady=3)
        else:
            e.pack(side="left", padx=(4,0))
        return e

    def _build(self):
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=0)
        self.grid_columnconfigure(1, weight=1)

        # ── LEFT PANEL: proper Canvas+Scrollbar, no bind_all ──────────
        lframe = tk.Frame(self, bg=CARD, width=410)
        lframe.grid(row=0, column=0, sticky="nsew")
        lframe.grid_propagate(False)
        lframe.grid_rowconfigure(0, weight=1)
        lframe.grid_columnconfigure(0, weight=1)

        canvas = tk.Canvas(lframe, bg=CARD, highlightthickness=0)
        vsb    = ttk.Scrollbar(lframe, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.grid(row=0, column=1, sticky="ns")
        canvas.grid(row=0, column=0, sticky="nsew")

        inner = tk.Frame(canvas, bg=CARD)
        win_id = canvas.create_window((0, 0), window=inner, anchor="nw")

        def _on_inner_cfg(e):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(win_id, width=canvas.winfo_width())
        def _on_canvas_cfg(e):
            canvas.itemconfig(win_id, width=canvas.winfo_width())
        # Use canvas.bind (not bind_all) so it only scrolls when mouse is over this canvas
        def _wheel(e):
            canvas.yview_scroll(int(-1*(e.delta/120)), "units")
        canvas.bind("<MouseWheel>", _wheel)
        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _wheel))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))
        inner.bind("<Configure>", _on_inner_cfg)
        canvas.bind("<Configure>", _on_canvas_cfg)

        # ── TITLE ─────────────────────────────────────────────────────
        tk.Label(inner, text="💰 Budget Manager", font=FH,
                 fg=ACCENT, bg=CARD).pack(anchor="w", padx=20, pady=(20,4))
        sep(inner, pady=4)

        # ── STEP 1: Income & Savings ───────────────────────────────────
        self._step_header(inner, "Step 1", "Income & Savings Goal", ACCENT)
        g1 = tk.Frame(inner, bg=CARD)
        g1.pack(fill="x", padx=20, pady=6)

        for r, (ltext, attr, default) in enumerate([
            ("Month (YYYY-MM):", "_ent_month",  datetime.datetime.now().strftime("%Y-%m")),
            ("Monthly Income ₹:", "_ent_income", "0"),
            ("Savings Goal ₹:",   "_ent_savings","0"),
        ]):
            tk.Label(g1, text=ltext, font=F, fg=DIM, bg=CARD,
                     width=18, anchor="w").grid(row=r, column=0, sticky="w", pady=4)
            ent = self._make_entry(g1, default, width=15, row=r, col=1)
            setattr(self, attr, ent)   # store direct widget reference

        sep(inner, pady=8)

        # ── STEP 2: Category Checkboxes ────────────────────────────────
        self._step_header(inner, "Step 2", "Select Your Categories", PURPLE)
        tk.Label(inner, text="  Tick the categories you spend money on:",
                 font=FS, fg=DIM, bg=CARD).pack(anchor="w", padx=20)

        cgrid = tk.Frame(inner, bg=CARD)
        cgrid.pack(fill="x", padx=20, pady=8)
        active = self.app.data["budget"].get("active_categories", ALL_CATEGORIES[:5])
        for i, cat in enumerate(ALL_CATEGORIES):
            v = tk.BooleanVar(value=(cat in active))
            self._cat_vars[cat] = v
            tk.Checkbutton(cgrid, text=cat, variable=v, font=F,
                           fg=TEXT, bg=CARD, activebackground=CARD,
                           selectcolor=INPUT,
                           command=self._on_cat_toggle
                           ).grid(row=i//2, column=i%2, sticky="w", padx=6, pady=3)

        sep(inner, pady=8)

        # ── STEP 3: Expense Amounts ────────────────────────────────────
        self._step_header(inner, "Step 3", "Expense Amounts", GREEN)
        btn(inner, "✨  Auto-Suggest Split  (50/30/20 Rule)",
            cmd=self._suggest_split, bg=GREEN, fg=BG,
            padx=14, pady=9).pack(fill="x", padx=20, pady=(4,6))
        tk.Label(inner,
                 text="  Amounts will be auto-filled based on your income.\n"
                      "  You can freely edit any field after.",
                 font=FS, fg=DIM, bg=CARD, justify="left").pack(anchor="w", padx=20)

        self._amt_frame = tk.Frame(inner, bg=CARD)
        self._amt_frame.pack(fill="x", padx=20, pady=8)
        self._rebuild_amt_fields()

        sep(inner, pady=8)

        # ── STEP 4: Custom Expenses ────────────────────────────────────
        self._step_header(inner, "Step 4", "Extra / Custom Expenses", ORANGE)
        add = tk.Frame(inner, bg=CARD)
        add.pack(fill="x", padx=20, pady=4)
        for r, (ltext, attr, w) in enumerate([
            ("Name:",      "_ent_cname", 14),
            ("Amount ₹:", "_ent_camt",  10),
        ]):
            tk.Label(add, text=ltext, font=FS, fg=DIM, bg=CARD,
                     width=10, anchor="w").grid(row=r, column=0, sticky="w", pady=2)
            ent = self._make_entry(add, "", width=w, row=r, col=1)
            setattr(self, attr, ent)

        tk.Label(add, text="Category:", font=FS, fg=DIM, bg=CARD,
                 width=10, anchor="w").grid(row=2, column=0, sticky="w")
        self._custom_cat_cb = ttk.Combobox(add, textvariable=self._v_ccat,
                                            values=ALL_CATEGORIES,
                                            state="readonly", width=13, font=FS)
        self._custom_cat_cb.grid(row=2, column=1, padx=(4,0), sticky="w", pady=2)

        btn(inner, "➕  Add Custom Expense",
            cmd=self._add_custom, bg=ORANGE, fg=BG,
            padx=14, pady=8).pack(fill="x", padx=20, pady=(8,4))

        tk.Label(inner, text="Custom Expense List:", font=FS, fg=DIM,
                 bg=CARD).pack(anchor="w", padx=20)
        self._custom_frame = tk.Frame(inner, bg=CARD)
        self._custom_frame.pack(fill="x", padx=20, pady=4)

        sep(inner, pady=8)

        # ── Action Buttons ─────────────────────────────────────────────
        btn(inner, "📊  Calculate Budget",
            cmd=self._calculate, bg=ACCENT, fg=BG,
            padx=18, pady=12).pack(fill="x", padx=20, pady=4)
        btn(inner, "💾  Save This Month",
            cmd=self._save_month, bg=PURPLE, fg=BG,
            padx=18, pady=10).pack(fill="x", padx=20, pady=4)
        btn(inner, "📤  Export CSV",
            cmd=self._export_csv, bg=INPUT, fg=DIM,
            padx=18, pady=8).pack(fill="x", padx=20, pady=4)
        tk.Label(inner, bg=CARD).pack(pady=20)

        # ── RIGHT PANEL ───────────────────────────────────────────────
        right = tk.Frame(self, bg=BG)
        right.grid(row=0, column=1, sticky="nsew", padx=(3,0))
        right.grid_rowconfigure(1, weight=1)
        right.grid_columnconfigure(0, weight=1)

        tk.Label(right, text="Budget Analysis", font=FH,
                 fg=BRIGHT, bg=BG).grid(row=0, column=0, sticky="w",
                                         padx=20, pady=(20,8))
        self._analysis = tk.Frame(right, bg=CARD)
        self._analysis.grid(row=1, column=0, sticky="nsew", padx=20, pady=(0,20))
        lbl(self._analysis,
            "\n\n  Complete Steps 1–4 on the left,\n"
            "  then click  📊 Calculate Budget.\n\n"
            "  New flow:\n"
            "  1️⃣  Enter income & savings goal\n"
            "  2️⃣  Tick your spending categories\n"
            "  3️⃣  Click 'Auto-Suggest Split'\n"
            "       → Edit any amount you like\n"
            "  4️⃣  Add any custom expenses\n"
            "  5️⃣  Calculate & see full analysis!\n",
            font=F, fg=DIM, bg=CARD, justify="left"
            ).pack(anchor="w", padx=24, pady=24)

    def _step_header(self, parent, num, title, color):
        f = tk.Frame(parent, bg=CARD)
        f.pack(fill="x", padx=20, pady=(10,2))
        tk.Label(f, text=num, font=("Segoe UI",9,"bold"), fg=BG,
                 bg=color, padx=6, pady=2).pack(side="left")
        tk.Label(f, text=f"  {title}", font=FM, fg=color, bg=CARD).pack(side="left")

    def _on_cat_toggle(self):
        """Rebuild amount fields when checkboxes change."""
        self._rebuild_amt_fields()

    def _rebuild_amt_fields(self, preset=None):
        """
        preset: optional dict {cat: amount_str} to pre-fill fields.
        Saves current values before destroying so edits aren't lost.
        """
        # Capture current values before destroying
        saved = {}
        for cat, ent in self._amt_entries.items():
            saved[cat] = ent.get()
        if preset:
            saved.update({c: str(v) for c, v in preset.items()})

        for w in self._amt_frame.winfo_children():
            w.destroy()
        self._amt_entries.clear()

        active = [c for c, v in self._cat_vars.items() if v.get()]
        if not active:
            tk.Label(self._amt_frame, text="  No categories selected yet.",
                     font=FS, fg=DIM, bg=CARD).pack(anchor="w")
            return

        CAT_COLORS = {
            "Rent": RED, "Food": GREEN, "Transport": YELLOW,
            "Shopping": PURPLE, "Entertainment": ACCENT,
            "Utilities": ORANGE, "Healthcare": AI_CLR,
            "Education": "#79c0ff", "Savings": GREEN, "Other": DIM
        }
        for cat in active:
            row = tk.Frame(self._amt_frame, bg=CARD)
            row.pack(fill="x", pady=2)
            dot_clr = CAT_COLORS.get(cat, DIM)
            tk.Label(row, text="●", font=("Segoe UI",10),
                     fg=dot_clr, bg=CARD).pack(side="left")
            tk.Label(row, text=f"  {cat}", font=F, fg=TEXT, bg=CARD,
                     width=14, anchor="w").pack(side="left")
            # Direct Entry widget — NO StringVar, read with .get() directly
            ent = tk.Entry(row, width=11, bg=INPUT, fg=BRIGHT,
                           insertbackground=BRIGHT, relief="flat", font=F, bd=4)
            ent.insert(0, saved.get(cat, "0"))
            ent.pack(side="left")
            tk.Label(row, text=" ₹", font=FS, fg=DIM, bg=CARD).pack(side="left")
            self._amt_entries[cat] = ent  # store direct widget reference

        if hasattr(self, "_custom_cat_cb"):
            self._custom_cat_cb["values"] = active + ["Other"]

    def _suggest_split(self):
        # Read directly from Entry widget, not StringVar
        try:
            income  = float(self._ent_income.get().strip())
        except (ValueError, AttributeError):
            messagebox.showerror("Invalid", "Enter a valid number for Monthly Income.")
            return
        try:
            savings = float(self._ent_savings.get().strip())
        except (ValueError, AttributeError):
            savings = 0

        if income <= 0:
            messagebox.showerror("Invalid", "Income must be greater than 0.")
            return
        active = [c for c, v in self._cat_vars.items() if v.get()]
        if not active:
            messagebox.showwarning("No Categories",
                "Please tick at least one category in Step 2 first.")
            return

        split = smart_split(income, savings, active)
        # Rebuild fields with suggested amounts pre-filled
        self._rebuild_amt_fields(preset={c: str(round(v, 2)) for c, v in split.items()})
        messagebox.showinfo("✨ Split Applied!",
            f"Auto-split for ₹{income-savings:,.0f} spendable budget.\n\n"
            "Based on 50/30/20 rule:\n"
            "  60% → Needs (Rent, Food, Utilities…)\n"
            "  30% → Wants (Shopping, Entertainment…)\n"
            "  10% → Flexibility\n\n"
            "You can edit any field freely now.")

    def _sync(self):
        """Read ALL values directly from Entry widget .get() — guaranteed accurate."""
        b = self.app.data["budget"]
        # Month
        b["month"] = self._ent_month.get().strip() if self._ent_month else \
                     datetime.datetime.now().strftime("%Y-%m")
        # Income
        try:    b["income"]       = float(self._ent_income.get().strip())
        except: b["income"]       = 0
        # Savings
        try:    b["savings_goal"] = float(self._ent_savings.get().strip())
        except: b["savings_goal"] = 0
        # Active categories
        b["active_categories"] = [c for c, v in self._cat_vars.items() if v.get()]
        # Category amounts — read directly from Entry widget
        b["expenses"] = {}
        for cat, ent in self._amt_entries.items():
            if self._cat_vars.get(cat, tk.BooleanVar(value=False)).get():
                try:    b["expenses"][cat] = float(ent.get().strip())
                except: b["expenses"][cat] = 0
        # Custom list
        b["custom_expenses"] = self._custom_list

    def _load(self):
        """Load saved data into widgets directly."""
        b = self.app.data["budget"]
        # Clear and set top fields directly
        if self._ent_month:
            self._ent_month.delete(0, "end")
            self._ent_month.insert(0, b.get("month", datetime.datetime.now().strftime("%Y-%m")))
        if self._ent_income:
            self._ent_income.delete(0, "end")
            self._ent_income.insert(0, str(b.get("income", 0)))
        if self._ent_savings:
            self._ent_savings.delete(0, "end")
            self._ent_savings.insert(0, str(b.get("savings_goal", 0)))
        # Checkboxes
        active = b.get("active_categories", ALL_CATEGORIES[:5])
        for cat, var in self._cat_vars.items():
            var.set(cat in active)
        # Rebuild amount fields with saved values
        saved_exps = b.get("expenses", {})
        self._rebuild_amt_fields(preset={c: str(v) for c, v in saved_exps.items()})
        # Custom list
        self._custom_list = list(b.get("custom_expenses", []))
        self._refresh_custom_frame()

    def _add_custom(self):
        name = self._ent_cname.get().strip() if self._ent_cname else ""
        cat  = self._v_ccat.get()
        try:
            raw = self._ent_camt.get().strip() if self._ent_camt else ""
            amt = float(raw)
            if amt < 0: raise ValueError
        except ValueError:
            messagebox.showerror("Invalid", "Enter a valid positive amount.")
            return
        if not name:
            messagebox.showerror("Invalid", "Enter an expense name.")
            return
        self._custom_list.append({"name": name, "amount": amt, "category": cat})
        if self._ent_cname:
            self._ent_cname.delete(0, "end")
        if self._ent_camt:
            self._ent_camt.delete(0, "end")
        self._refresh_custom_frame()

    def _refresh_custom_frame(self):
        for w in self._custom_frame.winfo_children():
            w.destroy()
        if not self._custom_list:
            tk.Label(self._custom_frame, text="  No custom expenses.",
                     font=FS, fg=DIM, bg=CARD).pack(anchor="w")
            return
        for i, exp in enumerate(self._custom_list):
            row = tk.Frame(self._custom_frame, bg=INPUT)
            row.pack(fill="x", pady=1)
            tk.Label(row, text=f"  {exp['name']} [{exp['category']}] ₹{exp['amount']:,.0f}",
                     font=FS, fg=TEXT, bg=INPUT).pack(side="left", padx=4, pady=4)
            tk.Button(row, text="✕", font=FS, fg=RED, bg=INPUT, relief="flat",
                      bd=0, cursor="hand2",
                      command=lambda i_=i: self._del_custom(i_)).pack(side="right", padx=6)

    def _del_custom(self, idx):
        self._custom_list.pop(idx)
        self._refresh_custom_frame()

    def _calculate(self):
        self._sync()
        a = calc_budget(self.app.data["budget"])
        self._render(a)

    def _render(self, a):
        for w in self._analysis.winfo_children():
            w.destroy()

        # Top stat row
        top = tk.Frame(self._analysis, bg=CARD)
        top.pack(fill="x", padx=14, pady=(14,8))
        for ltext, val, clr in [
            ("Income",       f"₹{a['income']:,.0f}",       ACCENT),
            ("Spendable",    f"₹{a['spendable']:,.0f}",    PURPLE),
            ("Total Spent",  f"₹{a['total_exp']:,.0f}",    ORANGE),
            ("Remaining",    f"₹{a['remaining']:,.0f}",    RED if a["over"] else GREEN),
        ]:
            f = tk.Frame(top, bg=INPUT, padx=10, pady=10)
            f.pack(side="left", expand=True, fill="both", padx=4)
            tk.Label(f, text=val, font=("Segoe UI",14,"bold"), fg=clr, bg=INPUT).pack(anchor="w")
            tk.Label(f, text=ltext, font=FS, fg=DIM, bg=INPUT).pack(anchor="w")

        # Status banner
        if a["over"]:
            excess = abs(a["remaining"])
            msg = (f"⚠  OVERSPENDING by ₹{excess:,.0f}\n"
                   f"   Top category: {a['top_cat']} (₹{a['top_amt']:,.0f})\n"
                   f"   Tip: Reduce {a['top_cat']} by 15–20% → save ₹{a['top_amt']*0.175:,.0f}")
            bg_, fg_ = "#2d0f0f", RED
        else:
            msg = (f"✅  Budget looks healthy! Remaining: ₹{a['remaining']:,.0f}\n"
                   f"   Top spend: {a['top_cat']} (₹{a['top_amt']:,.0f})")
            bg_, fg_ = "#0d2010", GREEN

        tk.Label(self._analysis, text=msg, font=FB, fg=fg_, bg=bg_,
                 padx=16, pady=10, justify="left", wraplength=560
                 ).pack(fill="x", padx=14, pady=4)

        tk.Frame(self._analysis, bg=BORDER, height=1).pack(fill="x", padx=14, pady=4)
        tk.Label(self._analysis, text="Category Breakdown", font=FM,
                 fg=ACCENT, bg=CARD).pack(anchor="w", padx=14, pady=(4,2))

        # Table
        tbl = tk.Frame(self._analysis, bg=CARD)
        tbl.pack(fill="x", padx=14, pady=4)
        for c, (hdr, w) in enumerate([("Category",14),("Amount ₹",10),("% Spend",10),("Bar",22)]):
            tk.Label(tbl, text=hdr, font=FB, fg=DIM, bg=INPUT,
                     width=w, anchor="w", padx=6, pady=3).grid(
                row=0, column=c, sticky="ew", padx=1, pady=1)

        total = max(a["total_exp"], 1)
        sorted_cats = sorted(a["cat_totals"].items(), key=lambda x: -x[1])
        for r, (cat, amt) in enumerate(sorted_cats, 1):
            pct  = amt / total * 100
            bar  = "█" * int(pct / 4)
            clr  = RED if cat == a["top_cat"] else TEXT
            rbg  = INPUT if r % 2 == 0 else CARD
            for c, (val, w) in enumerate([
                (cat, 14), (f"{amt:,.2f}", 10),
                (f"{pct:.1f}%", 10), (bar, 22)
            ]):
                tk.Label(tbl, text=val, font=FS, fg=clr, bg=rbg,
                         width=w, anchor="w", padx=6, pady=3).grid(
                    row=r, column=c, sticky="ew", padx=1, pady=1)

        # Suggestions
        if a["over"] or a["total_exp"] > 0:
            tk.Frame(self._analysis, bg=BORDER, height=1).pack(fill="x", padx=14, pady=6)
            tk.Label(self._analysis, text="💡 Smart Suggestions", font=FM,
                     fg=YELLOW, bg=CARD).pack(anchor="w", padx=14)
            suggestions = []
            if a["over"]:
                suggestions.append(f"• Reduce {a['top_cat']} spending by 15–20%")
                suggestions.append(f"• Current shortfall: ₹{abs(a['remaining']):,.0f}")
            if a["savings_goal"] == 0:
                suggestions.append("• Set a savings goal to better track your finances")
            spendable = a["spendable"]
            if spendable > 0 and a["total_exp"] / max(spendable, 1) > 0.9:
                suggestions.append("• You've spent >90% of spendable income — be cautious")
            if not suggestions:
                suggestions.append("• Finances look good! Keep maintaining this balance.")
            for s in suggestions:
                tk.Label(self._analysis, text=s, font=F, fg=TEXT, bg=CARD,
                          padx=14, pady=1).pack(anchor="w", padx=14)

        # Chart
        if HAS_MPL and a["total_exp"] > 0 and len(a["cat_totals"]) > 0:
            cf = tk.Frame(self._analysis, bg=CARD)
            cf.pack(fill="x", padx=14, pady=8)
            cats = [k for k, v in a["cat_totals"].items() if v > 0]
            vals = [a["cat_totals"][k] for k in cats]
            clrs = [RED if c == a["top_cat"] else "#4f8ef7" for c in cats]
            try:
                fig, ax = plt.subplots(figsize=(6, 2.6))
                fig.patch.set_facecolor(CARD)
                ax.set_facecolor(BG)
                ax.bar(cats, vals, color=clrs, edgecolor=CARD)
                ax.set_ylabel("₹", color=DIM, fontsize=8)
                ax.tick_params(colors=DIM, labelsize=7)
                for sp in ax.spines.values(): sp.set_edgecolor(BORDER)
                plt.xticks(rotation=30, ha="right")
                plt.tight_layout()
                canv = FigureCanvasTkAgg(fig, master=cf)
                canv.draw()
                canv.get_tk_widget().pack(fill="x")
                plt.close(fig)
            except Exception:
                pass

    def _save_month(self):
        self._sync()
        a = calc_budget(self.app.data["budget"])
        month = (self._ent_month.get().strip() if self._ent_month else "") or datetime.datetime.now().strftime("%Y-%m")
        self.app.data["history"][month] = {
            "income": a["income"], "savings_goal": a["savings_goal"],
            "total_expense": a["total_exp"], "remaining": a["remaining"],
            "highest_cat": a["top_cat"], "highest_amt": a["top_amt"],
            "cat_totals": a["cat_totals"],
            "saved_at": datetime.datetime.now().isoformat()
        }
        self.app.save()
        messagebox.showinfo("Saved", f"Budget for {month} saved to history ✅")

    def _export_csv(self):
        self._sync()
        a = calc_budget(self.app.data["budget"])
        month_str = self._ent_month.get().strip() if self._ent_month else "budget"
        path = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV","*.csv")],
            initialfile=f"Budget_{month_str}.csv")
        if not path: return
        with open(path, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["EduBudget Pro — Budget Export"])
            w.writerow(["Month", month_str])
            w.writerow(["Income", a["income"]])
            w.writerow(["Savings Goal", a["savings_goal"]])
            w.writerow(["Total Expenses", a["total_exp"]])
            w.writerow(["Remaining", a["remaining"]])
            w.writerow([])
            w.writerow(["Category", "Amount"])
            for cat, amt in a["cat_totals"].items():
                w.writerow([cat, amt])
            if self._custom_list:
                w.writerow([])
                w.writerow(["Custom Expenses"])
                w.writerow(["Name","Category","Amount"])
                for e in self._custom_list:
                    w.writerow([e["name"], e["category"], e["amount"]])
        messagebox.showinfo("Exported", f"Saved to:\n{path}")

    def refresh(self):
        self._load()


# ═══════════════════════════════════════════════════════════════════════
#  HISTORY PAGE
# ═══════════════════════════════════════════════════════════════════════
class HistoryPage(tk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent, bg=BG)
        self.app = app
        self._sel = None
        self._build()

    def _build(self):
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=0)
        self.grid_columnconfigure(1, weight=1)

        # Left: list
        left = tk.Frame(self, bg=CARD, width=250)
        left.grid(row=0, column=0, sticky="nsew")
        left.grid_propagate(False)

        tk.Label(left, text="📊 History", font=FH, fg=ACCENT, bg=CARD).pack(
            anchor="w", padx=20, pady=(20,4))
        sep(left, pady=4)
        tk.Label(left, text="Saved Months", font=FM, fg=DIM, bg=CARD).pack(
            anchor="w", padx=20, pady=(8,4))

        lf = tk.Frame(left, bg=CARD)
        lf.pack(fill="both", expand=True, padx=14, pady=4)
        lf.grid_rowconfigure(0, weight=1)
        lf.grid_columnconfigure(0, weight=1)
        self._lb = tk.Listbox(lf, bg=INPUT, fg=BRIGHT,
                               selectbackground=SEL, selectforeground="white",
                               font=F, relief="flat", bd=0, activestyle="none")
        self._lb.grid(row=0, column=0, sticky="nsew")
        vs = ttk.Scrollbar(lf, orient="vertical", command=self._lb.yview)
        vs.grid(row=0, column=1, sticky="ns")
        self._lb.config(yscrollcommand=vs.set)
        self._lb.bind("<<ListboxSelect>>", self._on_sel)

        sep(left, pady=8)
        tk.Label(left, text="Compare With:", font=FS, fg=DIM, bg=CARD).pack(
            anchor="w", padx=20, pady=(0,2))
        self._v_cmp = tk.StringVar(value="— None —")
        self._cmp_cb = ttk.Combobox(left, textvariable=self._v_cmp,
                                     state="readonly", width=17, font=FS)
        self._cmp_cb.pack(anchor="w", padx=20)
        btn(left, "🔍 Compare", cmd=self._compare, bg=PURPLE, padx=10, pady=7).pack(
            fill="x", padx=14, pady=6)
        btn(left, "🗑 Delete", cmd=self._delete, bg=RED, padx=10, pady=6).pack(
            fill="x", padx=14, pady=2)
        btn(left, "📤 Export All CSV", cmd=self._export_all,
            bg=INPUT, fg=DIM, padx=10, pady=6).pack(fill="x", padx=14, pady=2)
        tk.Label(left, bg=CARD).pack(pady=8)

        # Right: detail
        right = tk.Frame(self, bg=BG)
        right.grid(row=0, column=1, sticky="nsew", padx=(3,0))
        right.grid_rowconfigure(1, weight=1)
        right.grid_columnconfigure(0, weight=1)
        tk.Label(right, text="Detail View", font=FH, fg=BRIGHT, bg=BG).grid(
            row=0, column=0, sticky="w", padx=20, pady=(20,8))
        df = tk.Frame(right, bg=CARD)
        df.grid(row=1, column=0, sticky="nsew", padx=20, pady=(0,20))
        df.grid_rowconfigure(0, weight=1); df.grid_columnconfigure(0, weight=1)
        self._detail = scrolledtext.ScrolledText(
            df, bg=BG, fg=TEXT, font=("Courier New",10),
            insertbackground=TEXT, relief="flat", padx=14, pady=14,
            state="disabled", wrap="word")
        self._detail.grid(row=0, column=0, sticky="nsew")
        self.refresh()

    def _populate(self):
        self._lb.delete(0,"end")
        months = sorted(self.app.data.get("history",{}).keys(), reverse=True)
        for m in months:
            self._lb.insert("end", f"  {m}")
        self._cmp_cb["values"] = ["— None —"] + months
        self._v_cmp.set("— None —")

    def _on_sel(self, e=None):
        sel = self._lb.curselection()
        if not sel: return
        self._sel = self._lb.get(sel[0]).strip()
        self._show(self._sel)

    def _show(self, month, cmp_month=None):
        h = self.app.data.get("history",{})
        rec = h.get(month)
        if not rec: return
        W = 58
        L = []
        def p(s=""): L.append(s)
        def rule(c="─"): L.append(c*W)
        p("═"*W); p(f"  Budget Report — {month}".center(W)); p("═"*W)
        p(f"  Income:          ₹{rec.get('income',0):>12,.2f}")
        p(f"  Savings Goal:    ₹{rec.get('savings_goal',0):>12,.2f}")
        p(f"  Total Expenses:  ₹{rec.get('total_expense',0):>12,.2f}")
        p(f"  Remaining:       ₹{rec.get('remaining',0):>12,.2f}")
        p(f"  Top Category:     {rec.get('highest_cat','N/A')} (₹{rec.get('highest_amt',0):,.0f})")
        p(); rule(); p("  Category Breakdown:"); rule()
        for cat, amt in rec.get("cat_totals",{}).items():
            pct = (amt / rec.get("total_expense",1) * 100) if rec.get("total_expense",0) else 0
            p(f"  {cat:<18} ₹{amt:>10,.2f}  ({pct:.1f}%)")
        if cmp_month and cmp_month in h:
            prev = h[cmp_month]
            p(); p("═"*W); p(f"  COMPARISON: {month}  vs  {cmp_month}".center(W)); p("═"*W)
            delta = rec.get("total_expense",0) - prev.get("total_expense",0)
            pct_c = (delta / prev.get("total_expense",1) * 100) if prev.get("total_expense",0) else 0
            p(f"  Expense Change: ₹{abs(delta):,.0f}  {'↑ UP' if delta>0 else '↓ DOWN'}  ({abs(pct_c):.1f}%)")
            p(); rule(); p("  Category Changes:"); rule()
            all_c = set(list(rec.get("cat_totals",{}).keys()) +
                        list(prev.get("cat_totals",{}).keys()))
            changes = [(c,
                        rec.get("cat_totals",{}).get(c,0),
                        prev.get("cat_totals",{}).get(c,0),
                        rec.get("cat_totals",{}).get(c,0)-prev.get("cat_totals",{}).get(c,0))
                       for c in all_c]
            changes.sort(key=lambda x: abs(x[3]), reverse=True)
            big = changes[0][0] if changes else ""
            for cat, cur, prv, chg in changes:
                note = " ← BIGGEST CHANGE" if cat == big and abs(chg) > 0 else ""
                p(f"  {cat:<16} ₹{cur:>9,.0f}  ₹{prv:>9,.0f}  {'+' if chg>=0 else ''}₹{chg:>8,.0f}{note}")
        p(); p("═"*W)
        self._detail.config(state="normal")
        self._detail.delete("1.0","end")
        self._detail.insert("end", "\n".join(L))
        self._detail.config(state="disabled")

    def _compare(self):
        if not self._sel:
            messagebox.showinfo("Select","Click a month first."); return
        cmp = self._v_cmp.get()
        self._show(self._sel, cmp if cmp != "— None —" and cmp != self._sel else None)

    def _delete(self):
        if not self._sel: return
        if messagebox.askyesno("Delete", f"Delete record for {self._sel}?"):
            self.app.data["history"].pop(self._sel, None)
            self.app.save()
            self._sel = None
            self._detail.config(state="normal"); self._detail.delete("1.0","end"); self._detail.config(state="disabled")
            self._populate()

    def _export_all(self):
        h = self.app.data.get("history",{})
        if not h: messagebox.showinfo("Empty","No history to export."); return
        path = filedialog.asksaveasfilename(defaultextension=".csv",
                                             filetypes=[("CSV","*.csv")],
                                             initialfile="EduBudget_History.csv")
        if not path: return
        with open(path,"w",newline="",encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["Month","Income","Savings","Total Expense","Remaining","Top Category","Top Amount"])
            for month, rec in sorted(h.items()):
                w.writerow([month,rec.get("income",0),rec.get("savings_goal",0),
                             rec.get("total_expense",0),rec.get("remaining",0),
                             rec.get("highest_cat",""),rec.get("highest_amt",0)])
        messagebox.showinfo("Exported",f"Saved to:\n{path}")

    def refresh(self):
        self._populate()
        self._detail.config(state="normal"); self._detail.delete("1.0","end")
        h = self.app.data.get("history",{})
        self._detail.insert("end",
            f"\n  {len(h)} month(s) saved.\n\n  ← Click a month to view details.\n"
            if h else "\n\n  No history saved yet.\n\n  Go to Budget Manager → Save This Month.")
        self._detail.config(state="disabled")


# ═══════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    s = ttk.Style()
    try: s.theme_use("clam")
    except Exception: pass
    s.configure("TScrollbar", background=INPUT, troughcolor=BG,
                 bordercolor=BG, arrowcolor=DIM)
    s.configure("TCombobox", fieldbackground=INPUT, background=INPUT,
                 foreground=BRIGHT, arrowcolor=DIM,
                 selectbackground=SEL, selectforeground="white")
    s.map("TCombobox", fieldbackground=[("readonly", INPUT)])
    App().mainloop()